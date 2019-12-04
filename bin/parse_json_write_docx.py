#!/usr/bin/env python3.6


import json
import argparse
import os as os
import docx
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.shared import OxmlElement
from docx.oxml import ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import datetime 
import errno

# for the table widths as docx is fussy
def set_col_widths(table):
    widths = (Inches(2), Inches(2))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# for adding the page numbers (three functions)
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# needed just for linking to HIVdb website
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


if __name__ == "__main__":

    sample2subtype = {}
    parser = argparse.ArgumentParser()
    parser.add_argument('--json', required=True, help='input json file containing query results')
    parser.add_argument('--data', required=True, help='input text-tab delimited file with the dataset of patient data')
    parser.add_argument('--output', required=False, help='name of tab-delimited text file containing sample subtypes')
    parser.add_argument('--reports', required=False, action='store_true', help='if this flag is included, .docx reports will be produced for each sample')
    args = parser.parse_args()

    json_in = ""
    subtype_output = ""

    if args.json:
        json_in = args.json
    else:
        print ("json file cannot be read by parser")

    if args.data:
        data_in = args.data
    else:
        print ("the text-tab delimited file cannot be read by parser")


    if args.output:
        subtype_output = args.output
    else:
        subtype_output = "subtypes.txt"

    now = datetime.datetime.now()
    if now.month < 10:
        time_now = (str(now.day) + "-0" +str(now.month) + "-" + str(now.year))
    elif now.day <10:
        time_now = ("0" +str(now.day) + "-" +str(now.month) + "-" + str(now.year))
    elif now.month < 10 and now.day <10:
        time_now = ("0" +str(now.day) + "-0" +str(now.month) + "-" + str(now.year))
    else:
        time_now =(str(now.day) + "-" +str(now.month) + "-" + str(now.year))

    path = time_now + "_reports/"
    if args.reports:
        try:
            os.makedirs(os.path.dirname(path))
        except OSError as exc:
            if exc.errno != errno.EEXIST:
                raise
            pass

# genename dictionary
    genedict = {}
    genedict["PR"]="Protease"
    genedict["IN"]="Intergase"
    genedict["RT"]="Reverse Transcriptase"

# parse the text-tab delimited file with patient DataLossWarning
    patientdata = {}
    with open(data_in) as data_file:
      headerline=data_file.readline()
      headerline=headerline.rstrip()
      print(headerline)
      colnames = headerline.split('\t')
      #check_list = ['Your Sample ID','Our/Alternative ID','Sample collection date','Date of Birth','Initials or Name','Sex','Facility or clinic name','Sample Type','Viral Load','Viral load Date','Lab Request Date','Requesting Clinician','Email Requesting Clinician','Report prepared by','Report Date','Approved by']
      #assert colnames == check_list, "Columns are misaligned: {0} vs {1}".format(colnames, check_list)
      #if colnames[0]=="Your Sample ID" and colnames[1]=="Our/Alternative ID" and colnames[2]=="Sample collection date" and colnames[3]=="Date of Birth":#
      if headerline=="Your Sample ID	Our/Alternative ID	Sample collection date	Date of Birth	Initials or Name	Sex	Facility or clinic name	Sample Type	Viral Load	Viral load Date	Lab Request Date	Requesting Clinician	Email Requesting Clinician	Report prepared by	Report Date	Approved by":
        print("Properly formatted dataset")
        for line in data_file:
          line=line.rstrip()
          values=line.split('\t')
          for i in range(len(colnames)):
            if i < len(values):
              patientdata[(values[1],colnames[i])]=values[i]
            else:
              patientdata[(values[1],colnames[i])]=""
        
      else:
        print("The column labels in "+data_in+" are not as expected. Expecting:")
        print("Your Sample ID\tOur/Alternative ID\tSample collection date\tDate of Birth\tInitials or Name\tSex\tFacility or clinic name\tSample Type\tViral Load\tViral load Date\tLab Request Date\tRequesting Clinician\tEmail Requesting Clinician\tReport prepared by\tReport Date\tApproved by")
      

# parse the sierrapy json output for relevant information - subtype and DRMs
    with open(json_in) as json_file:
        data = json.load(json_file)

        for i in data:
            sample = i['inputSequence']['header']
            report_file_name = path +  sample + "_report.docx"
            print(patientdata[(sample,"Your Sample ID")])
            document = Document()
            #fn='template2.docx'
            #document = Document(fn)
            # sorting out the header in a table on the first page
            document.sections[0].different_first_page_header_footer=True
            header = document.sections[0].first_page_header
            htable=header.add_table(1, 3, Inches(6))
            htab_cells=htable.rows[0].cells
            htab_cells[1].width = Inches(5.6)
            ht0=htab_cells[1].add_paragraph('')
            ht0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kh=ht0.add_run()
            kh.add_picture('UVRIlogo_best.png', width=Inches(4))
            ht1=htab_cells[2].add_paragraph(sample)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
           
            
            footer = document.sections[0].first_page_footer
            htable=footer.add_table(1, 2, Inches(6))
            htab_cells=htable.rows[0].cells
            htab_cells[0].width = Inches(2)
            ht0=htab_cells[0].add_paragraph('report produced in collaboration with')
            ht0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ht1=htab_cells[1].add_paragraph()
            ht1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kh=ht1.add_run()
            kh.add_picture('CVRlogo.png', width=Inches(3.6))
            
            header = document.sections[0].header
            paragraph = header.paragraphs[0]
            paragraph.text = "\t\t"+sample
            
            # document formating
            paragraph_format = document.styles['Normal'].paragraph_format
            paragraph_format.line_spacing = 1
            paragraph_format.space_after = 1
            style = document.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(12)

            
            
            # sorting out the margins
            document.sections[0].top_margin = Cm(2)
            document.sections[0].bottom_margin = Cm(2)
            document.sections[0].left_margin = Cm(2.5)
            document.sections[0].right_margin = Cm(2.5)

            # add address
            document.add_paragraph("Molecular Virology Laboratory\nPO Box 49, Entebbe, Uganda.\nTel: (+256) (0)417 704 000\nEmail:\nWorld Health Organisation Designated National and Regional HIV Drug Resistance Laboratory\n")


            document.add_heading("Results Report", level=1)
            table = document.add_table(rows=11, cols=4)
            table.style = 'Table Grid'
            table.cell(0,0).width = Cm(4.5)
            rows_to_merge=[0,1,3,4,5,7,9]
            respecitve_labels=['Your Sample ID','Our/Alternative ID','Facility or clinic name','Sample Type','Sample collection date','Lab Request Date','Report prepared by']
            # merge three columns
            for r in range(len(rows_to_merge)):
              #row = table.rows[i]
              a = table.cell(rows_to_merge[r], 1)
              b = table.cell(rows_to_merge[r], 2)
              c = table.cell(rows_to_merge[r], 3)
              A = a.merge(b)
              C = A.merge(c)
              
              rowname=table.cell(rows_to_merge[r], 0)
              rowname_text=rowname.paragraphs[0].add_run(respecitve_labels[r])
              rowname_text.bold = True
              C.text = patientdata[(sample,respecitve_labels[r])]
            
            # rows that don't need merging
            patient_text=table.cell(2, 0).paragraphs[0].add_run("Patient Details:")
            patient_text.bold=True
            dob_text=table.cell(2, 1).paragraphs[0].add_run("Date of Birth:\n")
            dob_text.bold=True
            dob_text=table.cell(2, 1).paragraphs[0].add_run(patientdata[(sample,"Date of Birth")])            
            name_text=table.cell(2, 2).paragraphs[0].add_run("Initials (Given then Family Name):\n")
            name_text.bold=True
            name_text=table.cell(2, 2).paragraphs[0].add_run(patientdata[(sample,"Initials or Name")])  
            sex_text=table.cell(2, 3).paragraphs[0].add_run("Sex:\n")
            sex_text.bold=True
            sex_text=table.cell(2, 3).paragraphs[0].add_run(patientdata[(sample,"Sex")])  
 
            report_text=table.cell(10, 0).paragraphs[0].add_run("Report Date:")
            report_text.bold=True
            table.cell(10, 1).paragraphs[0].add_run(patientdata[(sample,"Report Date")])
            approved_text=table.cell(10, 2).paragraphs[0].add_run("Approved by:")
            approved_text.bold=True
            table.cell(10, 3).paragraphs[0].add_run(patientdata[(sample,"Approved by")])
              
            # dealing with the 3 column row7 and row 9
            a = table.cell(6, 2)
            b = table.cell(6, 3)
            A = a.merge(b)
            load_text=table.cell(6, 0).paragraphs[0].add_run("Viral load:")
            load_text.bold=True
            table.cell(6, 1).paragraphs[0].add_run(patientdata[(sample,"Viral Load")])
            loaddate_text=A.paragraphs[0].add_run("Date:")
            loaddate_text.bold=True
            loaddate_text=A.paragraphs[0].add_run(patientdata[(sample,"Viral load Date")])
            
            a = table.cell(8, 2)
            b = table.cell(8, 3)
            A = a.merge(b)
            load_text=table.cell(8, 0).paragraphs[0].add_run("Requesting Clinician:")
            load_text.bold=True
            table.cell(8, 1).paragraphs[0].add_run(patientdata[(sample,"Requesting Clinician")])
            loaddate_text=A.paragraphs[0].add_run("Email:")
            loaddate_text.bold=True
            loaddate_text=A.paragraphs[0].add_run(patientdata[(sample,"Email Requesting Clinician")])
                        
            document.add_heading("HIV Drug Resistance Genotype Report", level=1)
            p0=document.add_paragraph("Below are the results from the ")
            HIVbold=p0.add_run("HIVdb Program ")
            HIVbold.bold=True
            p0.add_run("drug resistance interpretation from Stanford University HIV Drug Resistance Database ")
            hyperlink = add_hyperlink(p0, 'http://hivdb.stanford.edu/','(http://hivdb.stanford.edu/). ',None,True)
            p0.add_run("For any queries or assistance interpreting these results please contact the MRC/UVRI Basic Science Virology Lab.")
            #p0.add_hyperlink(text='foobar', url='http://github.com')
            
            
            document.add_heading("Sequence Summary", level=1)
# gene name and codon information
            for j in i["alignedGeneSequences"]:
                start = j["firstAA"]
                end = j["lastAA"]
                gene = j["gene"]["name"]

                p1 = document.add_paragraph("Sequence includes " + genedict[gene] + " (" + gene +"): codons " + str(start) + "-" + str(end))
# subtype information
            document.add_heading("HIV Subtype Determination", level=1)
            subtype = i['subtypeText']
            subtype_id = str.split(subtype, "(")[0]
            subtype_message = ""
            if subtype == 'NA':
                subtype_message = "No subtype information for sample:\t" + sample
            else:
                sample2subtype[sample] = subtype
                subtype_message = "Subtype: "+subtype
            p2 = document.add_paragraph(subtype_message)

# Drug resistance information
            for j in i["drugResistance"]:
                currentGene = j["gene"]["name"]
                document.add_heading ("Drug Resistance Interpretation: " + currentGene + "\n", level=1)
                mutations_dict = {}
                for k in j["mutationsByTypes"]:
                    mutation_type = k["mutationType"]
                    mutation_list = []
                    for m in k["mutations"]:
                        for key in m:
                            mutation_list.append(m[key])
                    mutation_string = ""
                    if not mutation_list:
                        mutation_string = "None"
                    else:
                        mutation_string = ", ".join(mutation_list)
                    mutations_dict[mutation_type] = mutation_string


                scores_dict = {}
                numerical_scores_dict = {}
                drug_class = ""
                drug_name = ""
                drug_abbr = ""
                drug_fullname = ""
                drug_score = ""
                drug_text = ""
                drug_id = ""
                drug_value = ""
                m_info = {}


                if currentGene == 'RT':
                    N = {}
                    NN = {}
                    for key, value in mutations_dict.items():
                        if key == 'Other':
                            p3 = document.add_paragraph(key + " Mutations: " + value)
                        else:
                            p3 = document.add_paragraph(key + " Resistance Mutations: " + value)
                    for k in j["drugScores"]:
                        drug_class = k["drugClass"]["name"]
                        drug_name = k["drug"]["name"]
                        drug_abbr = k["drug"]["displayAbbr"]
                        drug_fullname = k["drug"]["fullName"]
                        drug_score = k["score"]
                        drug_text = k["text"]

                        drug_id = drug_fullname + " (" + drug_abbr + ") "
                        drug_value = [drug_score, drug_text]
                        if drug_class == "NRTI":
                            N[drug_id] = drug_value
                        elif drug_class =="NNRTI":
                            NN[drug_id] = drug_value

                        m_name = ""
                        m_type = ""
                        m_text = ""
                        if drug_score != 0.0:

                            for p in k["partialScores"]:
                                for key, value in p.items():
                                    if key =='mutations':
                                        for x, y in value[0].items():
                                            if x == 'text':
                                                m_name = y
                                            if x == 'comments':
                                                m_text = y[0]["text"]
                                            m_info[m_name] = m_text
                    document.add_heading("Nucleoside Reverse Transcriptase Inhibitors", level=2)
                    p4 = document.add_paragraph()
                    p4table = document.add_table(rows=len(N.keys()), cols=2)
                    r=0
                    for key, value in N.items():
                        load_text=p4table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        p4table.cell(r, 1).paragraphs[0].add_run(value[1])
                        r +=1
                    p4table=set_col_widths(p4table)
                        
                    document.add_heading("Non-Nucleoside Reverse Transcriptase Inhibitors", level=2)
                    p5 = document.add_paragraph()
                    p5table = document.add_table(rows=len(NN.keys()), cols=2)
                    r=0
                    for key, value in NN.items():
                        #p5.add_run(key).bold = True
                        #p5.add_run("\t" + value[1] + "\n")
                        load_text=p5table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        p5table.cell(r, 1).paragraphs[0].add_run(value[1])
                        r += 1
                    p5table=set_col_widths(p5table)

                    document.add_heading ("Mutation Scoring: " + currentGene + "\n", level=1)
                    p7 = document.add_paragraph("Nucleoside Reverse Transcriptase Inhibitors\n")
                    p7table = document.add_table(rows=len(N.keys()), cols=2)
                    r=0
                    for key, value in N.items():
                        #p7.add_run(key).bold = True
                        load_text=p7table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        score  = value[0]
                        if score != 0.0:
                            #p7.add_run("\t" + str(score) + "\n").bold=True
                            load_text=p7table.cell(r, 1).paragraphs[0].add_run(str(score))
                            load_text.bold = True
                        else:
                            #p7.add_run("\t" + str(score) + "\n")
                            p7table.cell(r, 1).paragraphs[0].add_run(str(score))
                        r +=1
                    p7table=set_col_widths(p7table)
                    
                    p8 = document.add_paragraph("\nNon-Nucleoside Reverse Transcriptase Inhibitors\n")
                    p8table = document.add_table(rows=len(NN.keys()), cols=2)
                    r=0
                    for key, value in NN.items():
                        #p8.add_run(key).bold = True
                        load_text=p8table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        score  = value[0]
                        if score != 0.0:
                            #p8.add_run("\t" + str(score) + "\n").bold=True
                            load_text=p8table.cell(r, 1).paragraphs[0].add_run(str(score))
                            load_text.bold = True
                        else:
                            #p8.add_run("\t" + str(score) + "\n")
                            p8table.cell(r, 1).paragraphs[0].add_run(str(score))
                        r +=1
                    p8table=set_col_widths(p8table) 
                        
                    if not m_info:
                        continue
                    else:
                        document.add_heading(currentGene + " Comments", level=2)
                        for key, value in m_info.items():
                            p9 = document.add_paragraph(value, style='List Bullet')


                elif currentGene =='PR' or currentGene == 'IN':
                    for key, value in mutations_dict.items():
                        if key == 'Other':
                            p3 = document.add_paragraph(drug_class + " " + key + " Mutations: " + value)
                        else:
                            p3 = document.add_paragraph(drug_class + " " + key + " Resistance Mutations: " + value)


                    for k in j["drugScores"]:
                        drug_class = k["drugClass"]["name"]
                        drug_name = k["drug"]["name"]
                        drug_abbr = k["drug"]["displayAbbr"]
                        drug_fullname = k["drug"]["fullName"]
                        drug_score = k["score"]
                        drug_text = k["text"]

                        drug_id = drug_fullname + " (" + drug_abbr + ") "
                        drug_value = [drug_score, drug_text]

                        scores_dict[drug_id] = drug_value
                        # drug scores and comments
                        m_name = ""
                        m_type = ""
                        m_text = ""
                        if drug_score != 0.0:

                            for p in k["partialScores"]:
                                for key, value in p.items():
                                    if key =='mutations':
                                        for x, y in value[0].items():
                                            if x == 'text':
                                                m_name = y
                                            if x == 'comments':
                                                m_text = y[0]["text"]
                                            m_info[m_name] = m_text


                    document.add_heading(drug_class, level=2)
                    p4 = document.add_paragraph()

# THIS IS PROBABLY WHERE I CAN PRINT OUT IN A TABLE
                    p4table = document.add_table(rows=len(scores_dict.keys()), cols=2)
                    r=0
                    for key, value in scores_dict.items():
                        #p4.add_run(key + "\t").bold = True
                        #p4.add_run(value[1] + "\n")
                        load_text=p4table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        p4table.cell(r, 1).paragraphs[0].add_run(value[1])
                        r += 1
                    p4table=set_col_widths(p4table) 

                    document.add_heading ("Mutation Scoring: " + currentGene + "\n", level=1)
                    p6 = document.add_paragraph()
                    p6table = document.add_table(rows=len(scores_dict.keys()), cols=2)
                    r=0
                    for key, value in scores_dict.items():
                        #p6.add_run(key).bold = True
                        load_text=p6table.cell(r, 0).paragraphs[0].add_run(key)
                        load_text.bold = True
                        score  = value[0]
                        if score != 0.0:
                            #p6.add_run("\t" + str(score) + "\n").bold=True
                            load_text=p6table.cell(r, 1).paragraphs[0].add_run(str(score))
                            load_text.bold = True
                        else:
                            #p6.add_run("\t" + str(score) + "\n")
                            p6table.cell(r, 1).paragraphs[0].add_run(str(score))
                        r += 1
                    p6table=set_col_widths(p6table)    
                    

                    if not m_info:
                        continue
                    else:
                        document.add_heading(currentGene + " Comments", level=2)
                        for key, value in m_info.items():
                            p7 = document.add_paragraph(value, style='List Bullet')
                else:
                    print("Gene is not PR or RT")
            add_page_number(document.sections[0].footer.paragraphs[0].add_run())
            if args.reports:
               # print ("Saving report: " + report_file_name)
                document.save(report_file_name)
    with open(subtype_output, "w+") as out:
        for i in sample2subtype:
            out.write(i + "\t" + sample2subtype[i] + "\n")


